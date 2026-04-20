"""Generate 6 test documents covering all route branches of Universal Auto-Router.

Each doc contains a unique "canary phrase" so we can verify end-to-end retrieval.
"""
import os
from pathlib import Path

OUT = Path("/tmp/test-docs")
OUT.mkdir(exist_ok=True)

CANARY = {
    "text_pdf":    "АГМАЙНД-КАНАРЕЙКА-ТЕКСТ-0451 коэффициент плотности редуктора 1.47",
    "scan_pdf":    "АГМАЙНД-КАНАРЕЙКА-СКАН-2077 допустимый износ втулки 0.8мм",
    "visual_pdf":  "АГМАЙНД-КАНАРЕЙКА-ВИЖУАЛ-3301 температура обмотки статора 115C",
    "docx":        "АГМАЙНД-КАНАРЕЙКА-ДОКС-4815 партномер SK-9922-B1 количество 12 шт",
    "xlsx":        "АГМАЙНД-КАНАРЕЙКА-ЭКСЕЛЬ-1618 кварцевый резонатор 32768 Hz тип HC-49",
    "png_image":   "АГМАЙНД-КАНАРЕЙКА-ПНГ-0707 серийный блок АБ-77 заводской номер 195847",
}


def gen_text_pdf():
    """Text-heavy PDF: technical manual style. ReportLab → real text layer."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # Register a font that supports Cyrillic
    try:
        pdfmetrics.registerFont(TTFont("DejaVu", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"))
        font = "DejaVu"
    except Exception:
        font = "Helvetica"

    path = OUT / "01_text_manual.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=A4)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("Body", fontName=font, fontSize=11, leading=15))
    styles.add(ParagraphStyle("H1", fontName=font, fontSize=18, leading=22, spaceAfter=12))
    styles.add(ParagraphStyle("H2", fontName=font, fontSize=14, leading=18, spaceAfter=8))

    content = []
    content.append(Paragraph("Руководство по эксплуатации механизма", styles["H1"]))
    content.append(Spacer(1, 12))
    for chap_n, chap_title in enumerate([
        "Общие положения",
        "Конструктивные особенности",
        "Порядок технического обслуживания",
        "Регламент проверки узлов",
    ], start=1):
        content.append(Paragraph(f"Раздел {chap_n}. {chap_title}", styles["H2"]))
        for _ in range(3):
            content.append(Paragraph(
                "При эксплуатации оборудования соблюдать регламент заводской документации. "
                "Все параметры должны находиться в установленных пределах согласно таблицам 1—3. "
                "Проверка производится не реже одного раза в квартал специалистом, имеющим допуск "
                "к работам с оборудованием данного класса.",
                styles["Body"]
            ))
            content.append(Spacer(1, 6))
        content.append(Paragraph(
            "Контрольное измерение: " + CANARY["text_pdf"] if chap_n == 2 else
            "Параметр находится в пределах нормы.",
            styles["Body"]
        ))
        content.append(Spacer(1, 10))
    # Pad to ~8 pages
    for _ in range(30):
        content.append(Paragraph(
            "Дополнительные технические характеристики приведены в приложении А. "
            "Контрольные значения параметров и допуски указаны в таблице 4. "
            "Ответственный инженер обязан фиксировать все отклонения в журнале регламентных работ. "
            "Результаты проверок хранятся 5 лет с момента внесения в журнал.",
            styles["Body"]
        ))
    doc.build(content)
    return path


def gen_scan_pdf():
    """Scanned PDF: text rendered to image → wrapped in PDF (no text layer)."""
    from PIL import Image, ImageDraw, ImageFont
    import fitz

    # Render text as image
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 20)
        font_h = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 28)
    except Exception:
        font = font_h = ImageFont.load_default()

    lines = [
        "АКТ № 442 от 14 марта 2026 г.",
        "",
        "Комиссия в составе:",
        "  Председатель — Иванов А.П.",
        "  Члены комиссии — Сидоров К.М., Петров В.И.",
        "",
        "Произвела осмотр подшипникового узла механизма и установила следующее:",
        "",
        "1. Износ наружной обоймы в пределах нормы.",
        "2. " + CANARY["scan_pdf"],
        "3. Смазочный материал соответствует паспорту.",
        "4. Температурный режим — рабочий.",
        "",
        "Заключение: механизм допущен к дальнейшей эксплуатации.",
        "",
        "Подпись председателя:  _________________",
    ]

    img = Image.new("RGB", (1600, 2200), "white")
    draw = ImageDraw.Draw(img)
    y = 80
    for i, line in enumerate(lines):
        f = font_h if i == 0 else font
        draw.text((120, y), line, fill="black", font=f)
        y += 42 if i == 0 else 34

    # Add "scan imperfections" — slight noise and rotation illusion
    import random
    random.seed(42)
    for _ in range(500):
        x, py = random.randint(0, 1599), random.randint(0, 2199)
        draw.point((x, py), fill=random.choice(["#e0e0e0", "#d0d0d0", "#f0f0f0"]))

    # Wrap as PDF (no text layer — pure image)
    path = OUT / "02_scanned_act.pdf"
    img_path = str(OUT / "_scan_tmp.png")
    img.save(img_path, "PNG", dpi=(150, 150))
    pdf = fitz.open()
    rect = fitz.Rect(0, 0, 1600, 2200)
    page = pdf.new_page(width=1600, height=2200)
    page.insert_image(rect, filename=img_path)
    pdf.save(str(path))
    pdf.close()
    os.remove(img_path)
    return path


def gen_visual_pdf():
    """Visual-heavy PDF: charts + diagrams, little text."""
    import matplotlib.pyplot as plt
    from matplotlib.backends.backend_pdf import PdfPages
    import numpy as np

    path = OUT / "03_visual_report.pdf"
    with PdfPages(str(path)) as pdf:
        # Page 1: title + big chart
        fig, ax = plt.subplots(figsize=(8.27, 11.69))  # A4
        ax.set_title("Квартальный отчёт по параметрам", fontsize=18)
        x = np.linspace(0, 10, 100)
        ax.plot(x, np.sin(x) * 100 + 105, label="Температура, °C")
        ax.plot(x, np.cos(x) * 20 + 50, label="Давление, кПа")
        ax.set_xlabel("Время, ч")
        ax.legend()
        ax.grid(True)
        ax.text(0.5, -0.1, CANARY["visual_pdf"], fontsize=10,
                transform=ax.transAxes, ha="center")
        pdf.savefig(fig); plt.close(fig)

        # Pages 2-4: multi-panel charts
        for pg in range(3):
            fig, axes = plt.subplots(2, 2, figsize=(8.27, 11.69))
            for ax in axes.flat:
                data = np.random.randn(50).cumsum() + pg*10
                ax.plot(data)
                ax.fill_between(range(50), data, alpha=0.3)
                ax.set_title(f"Канал {np.random.randint(1,100)}")
            fig.suptitle(f"Диагностические графики — лист {pg+2}", fontsize=14)
            pdf.savefig(fig); plt.close(fig)
    return path


def gen_docx():
    """Office format — should route to Doc Extractor."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("Спецификация комплектации", level=1)
    doc.add_paragraph("Настоящая спецификация содержит перечень деталей и компонентов, "
                      "входящих в состав изделия согласно техническому заданию.")

    doc.add_heading("1. Основные узлы", level=2)
    for p in [
        "Корпус алюминиевый с порошковым покрытием.",
        "Электронный блок управления модель EBC-2200.",
        "Силовой кабель с маркировкой в соответствии ГОСТ 31565-2012.",
    ]:
        doc.add_paragraph(p, style="List Bullet")

    doc.add_heading("2. Комплектующие", level=2)
    t = doc.add_table(rows=1, cols=3)
    hdr = t.rows[0].cells
    hdr[0].text = "Позиция"
    hdr[1].text = "Наименование"
    hdr[2].text = "Кол-во"
    for pos, name, qty in [
        ("1.1", "Корпус алюминиевый", "1 шт"),
        ("1.2", "Плата контроллера EBC-2200", "1 шт"),
        ("1.3", CANARY["docx"], "1 комп"),
        ("1.4", "Уплотнение резиновое кольцевое", "4 шт"),
    ]:
        row = t.add_row().cells
        row[0].text = pos
        row[1].text = name
        row[2].text = qty

    doc.add_heading("3. Условия хранения", level=2)
    doc.add_paragraph("Температура хранения от +5 до +35 °C при относительной влажности не выше 80%.")

    path = OUT / "04_spec.docx"
    doc.save(str(path))
    return path


def gen_xlsx():
    """XLSX with tabular data."""
    from openpyxl import Workbook
    from openpyxl.styles import Font

    wb = Workbook()
    # Sheet 1: parts catalog
    ws = wb.active
    ws.title = "Каталог"
    headers = ["Артикул", "Название", "Производитель", "Цена", "Остаток"]
    for i, h in enumerate(headers, start=1):
        ws.cell(row=1, column=i, value=h).font = Font(bold=True)
    rows = [
        ("ART-001", "Подшипник шариковый 6205-2RS", "SKF", 320, 150),
        ("ART-002", "Вал выходной редуктора", "Ростсельмаш", 4800, 12),
        ("ART-003", CANARY["xlsx"], "Epson Toyocom", 45, 2000),
        ("ART-004", "Реле твердотельное 40А", "Fotek", 890, 47),
        ("ART-005", "Кабель силовой КГ 3×2.5", "Севкабель", 180, 300),
    ]
    for r, row in enumerate(rows, start=2):
        for c, val in enumerate(row, start=1):
            ws.cell(row=r, column=c, value=val)

    # Sheet 2: price history
    ws2 = wb.create_sheet("История цен")
    ws2.append(["Артикул", "Дата", "Цена"])
    for d in ["2025-01", "2025-04", "2025-07", "2025-10", "2026-01", "2026-04"]:
        ws2.append(["ART-003", d, 40 + hash(d) % 10])

    path = OUT / "05_catalog.xlsx"
    wb.save(str(path))
    return path


def gen_png():
    """PNG image — routes to docling for OCR."""
    from PIL import Image, ImageDraw, ImageFont

    try:
        f_title = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 32)
        f_body = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 22)
    except Exception:
        f_title = f_body = ImageFont.load_default()

    img = Image.new("RGB", (1200, 800), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 40), "МАРКИРОВКА ИЗДЕЛИЯ", fill="black", font=f_title)
    d.rectangle((40, 30, 1160, 120), outline="black", width=3)

    lines = [
        "Изготовитель: ООО АгроПромСервис",
        "Модель: РК-7800",
        "Дата выпуска: март 2026",
        "",
        CANARY["png_image"],
        "",
        "ТУ 23.18.29.112-001-2024",
    ]
    y = 180
    for ln in lines:
        d.text((80, y), ln, fill="black", font=f_body)
        y += 40
    # Make it look scanned: slight noise
    import random
    random.seed(7)
    for _ in range(300):
        x, py = random.randint(0, 1199), random.randint(0, 799)
        d.point((x, py), fill=random.choice(["#eeeeee", "#dddddd"]))

    path = OUT / "06_label.png"
    img.save(str(path), "PNG")
    return path


def main():
    results = []
    for name, gen in [
        ("text_pdf", gen_text_pdf),
        ("scan_pdf", gen_scan_pdf),
        ("visual_pdf", gen_visual_pdf),
        ("docx", gen_docx),
        ("xlsx", gen_xlsx),
        ("png_image", gen_png),
    ]:
        p = gen()
        size = p.stat().st_size
        print(f"  {name:<12} {p.name:<30} {size/1024:.1f} KB  canary=«{CANARY[name][:45]}...»")
        results.append((name, p, CANARY[name]))

    # Save manifest for later retrieval verification
    import json
    with open(OUT / "manifest.json", "w") as f:
        json.dump({name: {"path": str(p), "canary": c} for name, p, c in results},
                  f, ensure_ascii=False, indent=2)
    print(f"\nManifest: {OUT}/manifest.json")
    return results


if __name__ == "__main__":
    main()
